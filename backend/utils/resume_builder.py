"""
Resume builder utility.
- Extracts text from PDF files.
- Uses LLM to parse resume text into structured sections.
- Builds .docx files from structured form data using Classic or Modern templates.
"""

import json
import os
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from ingestion.filter import _call_llm
from utils.logger import get_logger

logger = get_logger(__name__)


# ── PDF extraction ────────────────────────────────────────────────────────────

def extract_text_from_pdf(file_path: str) -> str:
    try:
        from pypdf import PdfReader
        reader = PdfReader(file_path)
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages)
    except Exception as e:
        logger.error("Error reading PDF %s: %s", file_path, e)
        return ""


# ── LLM extraction ────────────────────────────────────────────────────────────

_LEVEL_LABELS = ["", "Beginner", "Elementary", "Intermediate", "Advanced", "Expert"]

_EXTRACT_SYSTEM = """You are a resume parser. Extract structured information from the resume text.
Return ONLY a valid JSON object with these exact keys (no markdown, no explanation):
{
  "full_name": "",
  "email": "",
  "phone": "",
  "location": "",
  "linkedin": "",
  "github": "",
  "website": "",
  "summary": "",
  "skills": [
    {
      "category": "Category name (e.g. Programming Languages, Frameworks, Tools, Cloud, Soft Skills)",
      "skills": [
        {"name": "skill name", "level": 3}
      ]
    }
  ],
  "experience": [
    {
      "company": "",
      "title": "",
      "start_date": "",
      "end_date": "",
      "location": "",
      "bullets": ["bullet1", "bullet2"]
    }
  ],
  "education": [
    {
      "institution": "",
      "degree": "",
      "field": "",
      "start_date": "",
      "end_date": "",
      "gpa": ""
    }
  ],
  "achievements": ["achievement1"],
  "certifications": [
    {
      "name": "",
      "issuer": "",
      "date": ""
    }
  ],
  "projects": [
    {
      "name": "",
      "link": "",
      "tech_stack": ["tech1", "tech2"],
      "description": ""
    }
  ]
}
For skills: group them into meaningful categories. Infer proficiency level (1=Beginner, 2=Elementary, 3=Intermediate, 4=Advanced, 5=Expert) from context clues like years of use, job titles, or emphasis in the resume.
For projects: extract personal/side/academic projects if mentioned. tech_stack should be an array of technology names.
Extract as much information as possible. Use empty strings/arrays when information is absent."""


def extract_resume_sections_llm(text: str) -> dict:
    """Call LLM to parse resume text into structured sections."""
    user_msg = f"Resume text:\n{text[:8000]}"
    raw = _call_llm(_EXTRACT_SYSTEM, user_msg)

    # Strip markdown code fences if present
    raw = raw.strip()
    if raw.startswith("```"):
        parts = raw.split("```")
        raw = parts[1] if len(parts) > 1 else raw
        if raw.startswith("json"):
            raw = raw[4:]
    raw = raw.strip()

    try:
        return json.loads(raw)
    except Exception as e:
        logger.error("Failed to parse LLM extraction JSON: %s\nRaw: %s", e, raw[:500])
        return {}


# ── DOCX helpers ──────────────────────────────────────────────────────────────

def _add_bottom_border(paragraph, color_hex: str = "000000", size: int = 6):
    """Add a bottom border to a paragraph (horizontal rule effect)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_left_border(paragraph, color_hex: str = "4F46E5", size: int = 18):
    """Add a left border accent to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), color_hex)
    pBdr.append(left)
    pPr.append(pBdr)


def _set_para_spacing(paragraph, before: int = 0, after: int = 0, line: int = None):
    pPr = paragraph._p.get_or_add_pPr()
    pSp = OxmlElement("w:spacing")
    pSp.set(qn("w:before"), str(before))
    pSp.set(qn("w:after"), str(after))
    if line:
        pSp.set(qn("w:line"), str(line))
        pSp.set(qn("w:lineRule"), "auto")
    pPr.append(pSp)


def _normalize_list(value: Any) -> list:
    """Return value as a list, handling strings and None."""
    if not value:
        return []
    if isinstance(value, list):
        return value
    return [s.strip() for s in str(value).replace(",", "\n").split("\n") if s.strip()]


def _normalize_skill_categories(value: Any) -> list:
    """Normalise skills to list of {category, skills:[{name, level}]}.
    Handles: new categorised format, legacy flat string list, and plain string."""
    if not value:
        return []
    # Already categorised (list of dicts with 'category' key)
    if isinstance(value, list) and value and isinstance(value[0], dict) and "category" in value[0]:
        result = []
        for cat in value:
            skills = []
            for s in cat.get("skills", []):
                if isinstance(s, dict):
                    skills.append({"name": str(s.get("name", "")), "level": int(s.get("level", 3))})
                elif isinstance(s, str):
                    skills.append({"name": s, "level": 3})
            if skills:
                result.append({"category": str(cat.get("category", "Skills")), "skills": skills})
        return result
    # Legacy flat list of strings
    if isinstance(value, list):
        flat = [str(s).strip() for s in value if str(s).strip()]
        if flat:
            return [{"category": "Skills", "skills": [{"name": s, "level": 3} for s in flat]}]
    # Plain comma/newline string
    if isinstance(value, str):
        flat = [s.strip() for s in value.replace(",", "\n").split("\n") if s.strip()]
        if flat:
            return [{"category": "Skills", "skills": [{"name": s, "level": 3} for s in flat]}]
    return []


def _level_dots(level: int) -> str:
    """Return a 5-dot proficiency string, e.g. level 3 → '●●●○○'."""
    level = max(1, min(5, level))
    return "●" * level + "○" * (5 - level)


def _normalize_projects(value: Any) -> list:
    """Return list of {name, link, tech_stack:list, description} dicts."""
    if not isinstance(value, list):
        return []
    result = []
    for p in value:
        if not isinstance(p, dict):
            continue
        tech = p.get("tech_stack", [])
        if isinstance(tech, str):
            tech = [t.strip() for t in tech.split(",") if t.strip()]
        elif not isinstance(tech, list):
            tech = []
        result.append({
            "name": str(p.get("name", "")),
            "link": str(p.get("link", "")),
            "tech_stack": [str(t) for t in tech],
            "description": str(p.get("description", "")),
        })
    return result


def _normalize_bullets(value: Any) -> list:
    """Return bullets as list from either list or newline-delimited string."""
    if isinstance(value, list):
        return [b for b in value if b and str(b).strip()]
    if isinstance(value, str):
        return [b.strip().lstrip("-•* ") for b in value.split("\n") if b.strip()]
    return []


# ── Classic template ──────────────────────────────────────────────────────────

def build_resume_docx_classic(data: dict, output_path: str) -> str:
    doc = Document()

    # Page margins
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)

    # Remove default paragraph spacing
    style = doc.styles["Normal"]
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(2)
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    def section_heading(title: str):
        p = doc.add_paragraph()
        run = p.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        _add_bottom_border(p, "000000", 4)
        _set_para_spacing(p, before=120, after=40)
        return p

    # ── Name ──
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_name = p_name.add_run(data.get("full_name") or "Your Name")
    r_name.bold = True
    r_name.font.size = Pt(18)
    r_name.font.name = "Calibri"
    _set_para_spacing(p_name, before=0, after=40)

    # ── Contact ──
    contact_parts = [
        data.get("email"), data.get("phone"), data.get("location"),
        data.get("linkedin"), data.get("github"), data.get("website"),
    ]
    contact_str = "  |  ".join(x for x in contact_parts if x)
    if contact_str:
        p_contact = doc.add_paragraph()
        p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_contact = p_contact.add_run(contact_str)
        r_contact.font.size = Pt(9)
        r_contact.font.name = "Calibri"
        _set_para_spacing(p_contact, before=0, after=60)

    # ── Summary ──
    summary = data.get("summary", "").strip()
    if summary:
        section_heading("Professional Summary")
        p = doc.add_paragraph(summary)
        _set_para_spacing(p, before=20, after=40)

    # ── Experience ──
    experience = data.get("experience") or []
    if experience:
        section_heading("Experience")
        for exp in experience:
            company = exp.get("company", "")
            title = exp.get("title", "")
            start = exp.get("start_date", "")
            end = exp.get("end_date", "") or "Present"
            location = exp.get("location", "")
            bullets = _normalize_bullets(exp.get("bullets", []))

            # Company | dates
            p_co = doc.add_paragraph()
            r_co = p_co.add_run(company)
            r_co.bold = True
            r_co.font.size = Pt(10)
            date_str = f"  {start} – {end}" if start else ""
            if date_str:
                r_dt = p_co.add_run(date_str)
                r_dt.font.size = Pt(9)
                r_dt.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            _set_para_spacing(p_co, before=60, after=0)

            # Title | location
            p_ti = doc.add_paragraph()
            r_ti = p_ti.add_run(title)
            r_ti.italic = True
            r_ti.font.size = Pt(10)
            if location:
                r_loc = p_ti.add_run(f"  {location}")
                r_loc.font.size = Pt(9)
                r_loc.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            _set_para_spacing(p_ti, before=0, after=20)

            # Bullets
            for bullet in bullets:
                bp = doc.add_paragraph(style="List Bullet")
                bp.add_run(bullet).font.size = Pt(10)
                _set_para_spacing(bp, before=0, after=20)

    # ── Education ──
    education = data.get("education") or []
    if education:
        section_heading("Education")
        for edu in education:
            institution = edu.get("institution", "")
            degree = edu.get("degree", "")
            field = edu.get("field", "")
            start = edu.get("start_date", "")
            end = edu.get("end_date", "")
            gpa = edu.get("gpa", "")

            p_inst = doc.add_paragraph()
            r_inst = p_inst.add_run(institution)
            r_inst.bold = True
            r_inst.font.size = Pt(10)
            date_str = f"  {start} – {end}" if start else ""
            if date_str:
                r_dt = p_inst.add_run(date_str)
                r_dt.font.size = Pt(9)
                r_dt.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            _set_para_spacing(p_inst, before=60, after=0)

            deg_str = degree
            if field:
                deg_str += f" in {field}"
            if gpa:
                deg_str += f"  (GPA: {gpa})"
            if deg_str:
                p_deg = doc.add_paragraph(deg_str)
                _set_para_spacing(p_deg, before=0, after=20)

    # ── Skills ──
    skill_categories = _normalize_skill_categories(data.get("skills"))
    if skill_categories:
        section_heading("Skills")
        for cat in skill_categories:
            p_cat = doc.add_paragraph()
            r_cat = p_cat.add_run(cat["category"] + ":  ")
            r_cat.bold = True
            r_cat.font.size = Pt(10)
            skill_parts = []
            for sk in cat["skills"]:
                label = _LEVEL_LABELS[max(1, min(5, sk["level"]))]
                skill_parts.append(f"{sk['name']}  {_level_dots(sk['level'])}  {label}")
            p_cat.add_run("     ".join(skill_parts)).font.size = Pt(10)
            _set_para_spacing(p_cat, before=20, after=20)

    # ── Certifications ──
    certifications = data.get("certifications") or []
    if certifications:
        section_heading("Certifications")
        for cert in certifications:
            name = cert.get("name", "")
            issuer = cert.get("issuer", "")
            date = cert.get("date", "")
            p_cert = doc.add_paragraph()
            r_cert = p_cert.add_run(name)
            r_cert.bold = True
            r_cert.font.size = Pt(10)
            meta = "  |  ".join(x for x in [issuer, date] if x)
            if meta:
                r_meta = p_cert.add_run(f"  {meta}")
                r_meta.font.size = Pt(9)
                r_meta.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            _set_para_spacing(p_cert, before=40, after=20)

    # ── Projects ──
    projects = _normalize_projects(data.get("projects", []))
    if projects:
        section_heading("Projects")
        for proj in projects:
            p_proj = doc.add_paragraph()
            r_name = p_proj.add_run(proj["name"])
            r_name.bold = True
            r_name.font.size = Pt(10)
            if proj["link"]:
                r_link = p_proj.add_run(f"  |  {proj['link']}")
                r_link.font.size = Pt(9)
                r_link.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            _set_para_spacing(p_proj, before=60, after=0)
            if proj["tech_stack"]:
                p_stack = doc.add_paragraph()
                r_label = p_stack.add_run("Tech: ")
                r_label.italic = True
                r_label.font.size = Pt(9)
                p_stack.add_run("  •  ".join(proj["tech_stack"])).font.size = Pt(9)
                _set_para_spacing(p_stack, before=0, after=0)
            if proj["description"]:
                p_desc = doc.add_paragraph(proj["description"])
                p_desc.runs[0].font.size = Pt(10)
                _set_para_spacing(p_desc, before=20, after=20)

    # ── Achievements ──
    achievements = _normalize_bullets(data.get("achievements", []))
    if achievements:
        section_heading("Achievements")
        for ach in achievements:
            ap = doc.add_paragraph(style="List Bullet")
            ap.add_run(ach).font.size = Pt(10)
            _set_para_spacing(ap, before=0, after=20)

    doc.save(output_path)
    return output_path


# ── Elegant template (Georgia serif, gold accents) ────────────────────────────

_GOLD = "B45309"
_GOLD_RGB = (0xB4, 0x53, 0x09)


def build_resume_docx_elegant(data: dict, output_path: str) -> str:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)

    style = doc.styles["Normal"]
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(2)
    style.font.name = "Georgia"
    style.font.size = Pt(10)

    def _gold(run):
        run.font.color.rgb = RGBColor(*_GOLD_RGB)

    def section_heading(title: str):
        p = doc.add_paragraph()
        run = p.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Georgia"
        _gold(run)
        _add_bottom_border(p, _GOLD, 4)
        _set_para_spacing(p, before=140, after=50)
        return p

    # Name (centered)
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_name.add_run(data.get("full_name") or "Your Name")
    r.bold = True
    r.font.size = Pt(20)
    r.font.name = "Georgia"
    _set_para_spacing(p_name, before=0, after=40)

    # Contact
    contact_parts = [data.get("email"), data.get("phone"), data.get("location"),
                     data.get("linkedin"), data.get("github")]
    contact_str = "  •  ".join(x for x in contact_parts if x)
    if contact_str:
        p_c = doc.add_paragraph()
        p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_c = p_c.add_run(contact_str)
        r_c.font.size = Pt(9)
        r_c.font.name = "Georgia"
        _gold(r_c)
        _set_para_spacing(p_c, before=0, after=80)

    summary = (data.get("summary") or "").strip()
    if summary:
        section_heading("Profile")
        p = doc.add_paragraph(summary)
        p.runs[0].font.size = Pt(10)
        _set_para_spacing(p, before=20, after=40)

    experience = data.get("experience") or []
    if experience:
        section_heading("Experience")
        for exp in experience:
            bullets = _normalize_bullets(exp.get("bullets", []))
            p_co = doc.add_paragraph()
            r_co = p_co.add_run(exp.get("company", ""))
            r_co.bold = True
            r_co.font.size = Pt(10)
            date_str = f"  {exp.get('start_date','')} – {exp.get('end_date','') or 'Present'}"
            if exp.get("start_date"):
                r_dt = p_co.add_run(date_str)
                r_dt.font.size = Pt(9)
                r_dt.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            _set_para_spacing(p_co, before=60, after=0)

            p_ti = doc.add_paragraph()
            r_ti = p_ti.add_run(exp.get("title", ""))
            r_ti.italic = True
            r_ti.font.size = Pt(10)
            _gold(r_ti)
            _set_para_spacing(p_ti, before=0, after=20)

            for bullet in bullets:
                bp = doc.add_paragraph()
                r_dash = bp.add_run("–  ")
                r_dash.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
                bp.add_run(bullet).font.size = Pt(10)
                bp.paragraph_format.left_indent = Inches(0.2)
                _set_para_spacing(bp, before=0, after=20)

    education = data.get("education") or []
    if education:
        section_heading("Education")
        for edu in education:
            p_inst = doc.add_paragraph()
            r_inst = p_inst.add_run(edu.get("institution", ""))
            r_inst.bold = True
            r_inst.font.size = Pt(10)
            if edu.get("start_date"):
                r_dt = p_inst.add_run(f"  {edu.get('start_date','')} – {edu.get('end_date','')}")
                r_dt.font.size = Pt(9)
                r_dt.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            _set_para_spacing(p_inst, before=60, after=0)
            deg_str = edu.get("degree", "")
            if edu.get("field"):
                deg_str += f" in {edu.get('field')}"
            if edu.get("gpa"):
                deg_str += f"  •  GPA: {edu.get('gpa')}"
            if deg_str:
                p_deg = doc.add_paragraph(deg_str)
                p_deg.runs[0].italic = True
                _set_para_spacing(p_deg, before=0, after=20)

    skill_categories = _normalize_skill_categories(data.get("skills"))
    if skill_categories:
        section_heading("Skills")
        for cat in skill_categories:
            p_cat = doc.add_paragraph()
            r_cat = p_cat.add_run(cat["category"] + ": ")
            r_cat.bold = True
            _gold(r_cat)
            p_cat.add_run("  ".join(sk["name"] for sk in cat["skills"]))
            _set_para_spacing(p_cat, before=20, after=20)

    achievements = _normalize_bullets(data.get("achievements", []))
    if achievements:
        section_heading("Achievements")
        for ach in achievements:
            ap = doc.add_paragraph()
            ap.add_run("–  ").font.color.rgb = RGBColor(*_GOLD_RGB)
            ap.add_run(ach).font.size = Pt(10)
            ap.paragraph_format.left_indent = Inches(0.2)
            _set_para_spacing(ap, before=0, after=20)

    certifications = data.get("certifications") or []
    if certifications:
        section_heading("Certifications")
        for cert in certifications:
            p_cert = doc.add_paragraph()
            r_cert = p_cert.add_run(cert.get("name", ""))
            r_cert.bold = True
            meta = "  •  ".join(x for x in [cert.get("issuer"), cert.get("date")] if x)
            if meta:
                r_meta = p_cert.add_run(f"  {meta}")
                r_meta.font.size = Pt(9)
                r_meta.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            _set_para_spacing(p_cert, before=40, after=20)

    doc.save(output_path)
    return output_path


# ── Minimal template (clean sans-serif, no borders) ───────────────────────────

def build_resume_docx_minimal(data: dict, output_path: str) -> str:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.7)
    sec.bottom_margin = Inches(0.7)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)

    style = doc.styles["Normal"]
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(2)
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    _LABEL = RGBColor(0x9C, 0xA3, 0xAF)  # gray-400

    def section_heading(title: str):
        p = doc.add_paragraph()
        run = p.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = "Calibri"
        run.font.color.rgb = _LABEL
        _set_para_spacing(p, before=160, after=40, line=276)
        return p

    p_name = doc.add_paragraph()
    r_name = p_name.add_run(data.get("full_name") or "Your Name")
    r_name.font.size = Pt(20)
    r_name.font.name = "Calibri"
    _set_para_spacing(p_name, before=0, after=20)

    contact_parts = [data.get("email"), data.get("phone"), data.get("location")]
    contact_str = "   ·   ".join(x for x in contact_parts if x)
    if contact_str:
        p_c = doc.add_paragraph()
        r_c = p_c.add_run(contact_str)
        r_c.font.size = Pt(9)
        r_c.font.color.rgb = _LABEL
        _set_para_spacing(p_c, before=0, after=80)

    summary = (data.get("summary") or "").strip()
    if summary:
        section_heading("About")
        p = doc.add_paragraph(summary)
        _set_para_spacing(p, before=20, after=40)

    experience = data.get("experience") or []
    if experience:
        section_heading("Experience")
        for exp in experience:
            bullets = _normalize_bullets(exp.get("bullets", []))
            p_row = doc.add_paragraph()
            r_co = p_row.add_run(exp.get("company", ""))
            r_co.bold = True
            r_co.font.size = Pt(10)
            date_str = f"  {exp.get('start_date','')} – {exp.get('end_date','') or 'Present'}"
            if exp.get("start_date"):
                r_dt = p_row.add_run(date_str)
                r_dt.font.size = Pt(9)
                r_dt.font.color.rgb = _LABEL
            _set_para_spacing(p_row, before=60, after=0)

            p_ti = doc.add_paragraph()
            r_ti = p_ti.add_run(exp.get("title", ""))
            r_ti.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            _set_para_spacing(p_ti, before=0, after=20)

            for bullet in bullets:
                bp = doc.add_paragraph()
                r_dot = bp.add_run("·  ")
                r_dot.font.color.rgb = _LABEL
                bp.add_run(bullet).font.size = Pt(10)
                bp.paragraph_format.left_indent = Inches(0.15)
                _set_para_spacing(bp, before=0, after=20)

    education = data.get("education") or []
    if education:
        section_heading("Education")
        for edu in education:
            p_inst = doc.add_paragraph()
            r_inst = p_inst.add_run(edu.get("institution", ""))
            r_inst.bold = True
            if edu.get("start_date"):
                r_dt = p_inst.add_run(f"  {edu.get('start_date','')} – {edu.get('end_date','')}")
                r_dt.font.size = Pt(9)
                r_dt.font.color.rgb = _LABEL
            _set_para_spacing(p_inst, before=60, after=0)
            deg_str = edu.get("degree", "")
            if edu.get("field"):
                deg_str += f" in {edu.get('field')}"
            if deg_str:
                p_deg = doc.add_paragraph(deg_str)
                p_deg.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                _set_para_spacing(p_deg, before=0, after=20)

    skill_categories = _normalize_skill_categories(data.get("skills"))
    if skill_categories:
        section_heading("Skills")
        for cat in skill_categories:
            p_cat = doc.add_paragraph()
            r_cat = p_cat.add_run(cat["category"] + "  ")
            r_cat.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            p_cat.add_run("  ·  ".join(sk["name"] for sk in cat["skills"]))
            _set_para_spacing(p_cat, before=20, after=20)

    achievements = _normalize_bullets(data.get("achievements", []))
    if achievements:
        section_heading("Achievements")
        for ach in achievements:
            ap = doc.add_paragraph()
            ap.add_run("·  ").font.color.rgb = _LABEL
            ap.add_run(ach).font.size = Pt(10)
            ap.paragraph_format.left_indent = Inches(0.15)
            _set_para_spacing(ap, before=0, after=20)

    certifications = data.get("certifications") or []
    if certifications:
        section_heading("Certifications")
        for cert in certifications:
            p_cert = doc.add_paragraph()
            r_cert = p_cert.add_run(cert.get("name", ""))
            r_cert.bold = True
            meta = "  ·  ".join(x for x in [cert.get("issuer"), cert.get("date")] if x)
            if meta:
                r_meta = p_cert.add_run(f"  {meta}")
                r_meta.font.size = Pt(9)
                r_meta.font.color.rgb = _LABEL
            _set_para_spacing(p_cert, before=40, after=20)

    doc.save(output_path)
    return output_path


# ── Slate template (Verdana, teal accent) ─────────────────────────────────────

_TEAL = "0F766E"
_TEAL_RGB = (0x0F, 0x76, 0x6E)


def build_resume_docx_slate(data: dict, output_path: str) -> str:
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)

    style = doc.styles["Normal"]
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(2)
    style.font.name = "Verdana"
    style.font.size = Pt(9)

    def _teal(run):
        run.font.color.rgb = RGBColor(*_TEAL_RGB)

    def section_heading(title: str):
        p = doc.add_paragraph()
        run = p.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = "Verdana"
        _teal(run)
        _add_bottom_border(p, _TEAL, 6)
        _set_para_spacing(p, before=140, after=50)
        return p

    # Name with left border accent
    p_name = doc.add_paragraph()
    r_name = p_name.add_run(data.get("full_name") or "Your Name")
    r_name.bold = True
    r_name.font.size = Pt(20)
    r_name.font.name = "Verdana"
    r_name.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    _add_left_border(p_name, _TEAL, 24)
    _set_para_spacing(p_name, before=0, after=30)

    contact_parts = [data.get("email"), data.get("phone"), data.get("location"),
                     data.get("linkedin"), data.get("github")]
    contact_str = "  |  ".join(x for x in contact_parts if x)
    if contact_str:
        p_c = doc.add_paragraph()
        r_c = p_c.add_run(contact_str)
        r_c.font.size = Pt(8)
        r_c.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        _set_para_spacing(p_c, before=0, after=80)

    summary = (data.get("summary") or "").strip()
    if summary:
        section_heading("Summary")
        p = doc.add_paragraph(summary)
        _set_para_spacing(p, before=20, after=40)

    experience = data.get("experience") or []
    if experience:
        section_heading("Experience")
        for exp in experience:
            bullets = _normalize_bullets(exp.get("bullets", []))
            p_co = doc.add_paragraph()
            r_co = p_co.add_run(exp.get("company", ""))
            r_co.bold = True
            r_co.font.size = Pt(10)
            if exp.get("start_date"):
                r_dt = p_co.add_run(f"  {exp.get('start_date','')} – {exp.get('end_date','') or 'Present'}")
                r_dt.font.size = Pt(8)
                r_dt.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
            _set_para_spacing(p_co, before=80, after=0)

            p_ti = doc.add_paragraph()
            r_ti = p_ti.add_run(exp.get("title", ""))
            r_ti.italic = True
            r_ti.font.size = Pt(9)
            _teal(r_ti)
            _set_para_spacing(p_ti, before=0, after=20)

            for bullet in bullets:
                bp = doc.add_paragraph()
                r_dot = bp.add_run("◆  ")
                _teal(r_dot)
                r_dot.font.size = Pt(8)
                bp.add_run(bullet).font.size = Pt(9)
                bp.paragraph_format.left_indent = Inches(0.15)
                _set_para_spacing(bp, before=0, after=20)

    education = data.get("education") or []
    if education:
        section_heading("Education")
        for edu in education:
            p_inst = doc.add_paragraph()
            r_inst = p_inst.add_run(edu.get("institution", ""))
            r_inst.bold = True
            r_inst.font.size = Pt(10)
            if edu.get("start_date"):
                r_dt = p_inst.add_run(f"  {edu.get('start_date','')} – {edu.get('end_date','')}")
                r_dt.font.size = Pt(8)
                r_dt.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
            _set_para_spacing(p_inst, before=60, after=0)
            deg_str = edu.get("degree", "")
            if edu.get("field"):
                deg_str += f" in {edu.get('field')}"
            if deg_str:
                p_deg = doc.add_paragraph(deg_str)
                p_deg.runs[0].italic = True
                _set_para_spacing(p_deg, before=0, after=20)

    skill_categories = _normalize_skill_categories(data.get("skills"))
    if skill_categories:
        section_heading("Skills")
        for cat in skill_categories:
            p_cat = doc.add_paragraph()
            r_cat = p_cat.add_run(cat["category"] + ":  ")
            r_cat.bold = True
            _teal(r_cat)
            for i, sk in enumerate(cat["skills"]):
                if i > 0:
                    p_cat.add_run("     ")
                r_sk = p_cat.add_run(sk["name"] + "  ")
                r_sk.font.size = Pt(9)
                r_dots = p_cat.add_run(_level_dots(sk["level"]))
                r_dots.font.size = Pt(8)
                _teal(r_dots)
            _set_para_spacing(p_cat, before=20, after=20)

    achievements = _normalize_bullets(data.get("achievements", []))
    if achievements:
        section_heading("Achievements")
        for ach in achievements:
            ap = doc.add_paragraph()
            r_dot = ap.add_run("◆  ")
            _teal(r_dot)
            r_dot.font.size = Pt(8)
            ap.add_run(ach).font.size = Pt(9)
            ap.paragraph_format.left_indent = Inches(0.15)
            _set_para_spacing(ap, before=0, after=20)

    certifications = data.get("certifications") or []
    if certifications:
        section_heading("Certifications")
        for cert in certifications:
            p_cert = doc.add_paragraph()
            r_cert = p_cert.add_run(cert.get("name", ""))
            r_cert.bold = True
            meta = "  ·  ".join(x for x in [cert.get("issuer"), cert.get("date")] if x)
            if meta:
                r_meta = p_cert.add_run(f"  {meta}")
                r_meta.font.size = Pt(8)
                r_meta.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
            _set_para_spacing(p_cert, before=40, after=20)

    doc.save(output_path)
    return output_path


# ── Executive template (navy header band) ─────────────────────────────────────

_NAVY = "1E293B"
_NAVY_RGB = (0x1E, 0x29, 0x3B)


def build_resume_docx_executive(data: dict, output_path: str) -> str:
    from docx.oxml.ns import qn as _qn
    from lxml import etree

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0)       # header band goes to very top
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0)
    sec.right_margin = Inches(0)

    style = doc.styles["Normal"]
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(2)
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    def _navy(run):
        run.font.color.rgb = RGBColor(*_NAVY_RGB)

    # ── Navy header band (name + contact) ──
    p_band = doc.add_paragraph()
    p_band.paragraph_format.left_indent = Inches(0.8)
    p_band.paragraph_format.right_indent = Inches(0.8)
    _set_para_spacing(p_band, before=240, after=0)
    # Shade the paragraph background navy
    pPr = p_band._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), _NAVY)
    pPr.append(shd)
    r_name = p_band.add_run(data.get("full_name") or "Your Name")
    r_name.bold = True
    r_name.font.size = Pt(22)
    r_name.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    contact_parts = [data.get("email"), data.get("phone"), data.get("location"),
                     data.get("linkedin"), data.get("github")]
    contact_str = "  |  ".join(x for x in contact_parts if x)
    if contact_str:
        p_c = doc.add_paragraph()
        p_c.paragraph_format.left_indent = Inches(0.8)
        p_c.paragraph_format.right_indent = Inches(0.8)
        _set_para_spacing(p_c, before=0, after=160)
        pPr2 = p_c._p.get_or_add_pPr()
        shd2 = OxmlElement("w:shd")
        shd2.set(qn("w:val"), "clear")
        shd2.set(qn("w:color"), "auto")
        shd2.set(qn("w:fill"), _NAVY)
        pPr2.append(shd2)
        r_c = p_c.add_run(contact_str)
        r_c.font.size = Pt(9)
        r_c.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    def section_heading(title: str):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.8)
        p.paragraph_format.right_indent = Inches(0.8)
        run = p.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(*_NAVY_RGB)
        _add_bottom_border(p, _NAVY, 6)
        _set_para_spacing(p, before=140, after=50)
        return p

    def body_para(text=""):
        p = doc.add_paragraph(text)
        p.paragraph_format.left_indent = Inches(0.8)
        p.paragraph_format.right_indent = Inches(0.8)
        return p

    summary = (data.get("summary") or "").strip()
    if summary:
        section_heading("Executive Summary")
        p = body_para(summary)
        _set_para_spacing(p, before=20, after=40)

    experience = data.get("experience") or []
    if experience:
        section_heading("Professional Experience")
        for exp in experience:
            bullets = _normalize_bullets(exp.get("bullets", []))
            p_co = body_para()
            r_co = p_co.add_run(exp.get("company", ""))
            r_co.bold = True
            r_co.font.size = Pt(11)
            if exp.get("start_date"):
                r_dt = p_co.add_run(f"  {exp.get('start_date','')} – {exp.get('end_date','') or 'Present'}")
                r_dt.font.size = Pt(9)
                r_dt.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
            _set_para_spacing(p_co, before=80, after=0)

            p_ti = body_para()
            r_ti = p_ti.add_run(exp.get("title", ""))
            r_ti.italic = True
            r_ti.font.size = Pt(10)
            r_ti.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
            _set_para_spacing(p_ti, before=0, after=20)

            for bullet in bullets:
                bp = body_para()
                bp.add_run("•  ").font.size = Pt(9)
                bp.add_run(bullet).font.size = Pt(10)
                bp.paragraph_format.left_indent = Inches(0.95)
                _set_para_spacing(bp, before=0, after=20)

    education = data.get("education") or []
    if education:
        section_heading("Education")
        for edu in education:
            p_inst = body_para()
            r_inst = p_inst.add_run(edu.get("institution", ""))
            r_inst.bold = True
            r_inst.font.size = Pt(10)
            if edu.get("start_date"):
                r_dt = p_inst.add_run(f"  {edu.get('start_date','')} – {edu.get('end_date','')}")
                r_dt.font.size = Pt(9)
                r_dt.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
            _set_para_spacing(p_inst, before=60, after=0)
            deg_str = edu.get("degree", "")
            if edu.get("field"):
                deg_str += f" in {edu.get('field')}"
            if edu.get("gpa"):
                deg_str += f"  •  GPA: {edu.get('gpa')}"
            if deg_str:
                p_deg = body_para(deg_str)
                p_deg.runs[0].italic = True
                _set_para_spacing(p_deg, before=0, after=20)

    skill_categories = _normalize_skill_categories(data.get("skills"))
    if skill_categories:
        section_heading("Core Competencies")
        for cat in skill_categories:
            p_cat = body_para()
            r_cat = p_cat.add_run(cat["category"] + ": ")
            r_cat.bold = True
            r_cat.font.size = Pt(10)
            p_cat.add_run("  ".join(sk["name"] for sk in cat["skills"])).font.size = Pt(10)
            _set_para_spacing(p_cat, before=20, after=20)

    achievements = _normalize_bullets(data.get("achievements", []))
    if achievements:
        section_heading("Achievements")
        for ach in achievements:
            ap = body_para()
            ap.add_run("•  ").font.size = Pt(9)
            ap.add_run(ach).font.size = Pt(10)
            ap.paragraph_format.left_indent = Inches(0.95)
            _set_para_spacing(ap, before=0, after=20)

    certifications = data.get("certifications") or []
    if certifications:
        section_heading("Certifications")
        for cert in certifications:
            p_cert = body_para()
            r_cert = p_cert.add_run(cert.get("name", ""))
            r_cert.bold = True
            meta = "  ·  ".join(x for x in [cert.get("issuer"), cert.get("date")] if x)
            if meta:
                r_meta = p_cert.add_run(f"  {meta}")
                r_meta.font.size = Pt(9)
                r_meta.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
            _set_para_spacing(p_cert, before=40, after=20)

    doc.save(output_path)
    return output_path


# ── Modern template ───────────────────────────────────────────────────────────

_ACCENT = "4F46E5"   # indigo
_DARK   = "1E1B4B"   # deep indigo for name
_GRAY   = "6B7280"   # muted gray for meta


def build_resume_docx_modern(data: dict, output_path: str) -> str:
    doc = Document()

    sec = doc.sections[0]
    sec.top_margin = Inches(0.75)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.9)
    sec.right_margin = Inches(0.9)

    style = doc.styles["Normal"]
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(2)
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    def section_heading(title: str):
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(
            int(_ACCENT[:2], 16), int(_ACCENT[2:4], 16), int(_ACCENT[4:], 16)
        )
        _add_left_border(p, _ACCENT, 18)
        _set_para_spacing(p, before=140, after=40)
        return p

    def _gray_run(paragraph, text: str):
        r = paragraph.add_run(text)
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(
            int(_GRAY[:2], 16), int(_GRAY[2:4], 16), int(_GRAY[4:], 16)
        )
        r.font.name = "Calibri"
        return r

    # ── Name ──
    p_name = doc.add_paragraph()
    r_name = p_name.add_run(data.get("full_name") or "Your Name")
    r_name.bold = True
    r_name.font.size = Pt(22)
    r_name.font.name = "Calibri"
    r_name.font.color.rgb = RGBColor(
        int(_DARK[:2], 16), int(_DARK[2:4], 16), int(_DARK[4:], 16)
    )
    _set_para_spacing(p_name, before=0, after=30)

    # ── Contact ──
    contact_parts = [
        data.get("email"), data.get("phone"), data.get("location"),
        data.get("linkedin"), data.get("github"), data.get("website"),
    ]
    contact_str = "   ·   ".join(x for x in contact_parts if x)
    if contact_str:
        p_contact = doc.add_paragraph()
        _gray_run(p_contact, contact_str)
        _set_para_spacing(p_contact, before=0, after=80)

    # ── Summary ──
    summary = (data.get("summary") or "").strip()
    if summary:
        section_heading("Summary")
        p = doc.add_paragraph(summary)
        p.runs[0].font.size = Pt(10)
        _set_para_spacing(p, before=20, after=40)

    # ── Experience ──
    experience = data.get("experience") or []
    if experience:
        section_heading("Experience")
        for exp in experience:
            company = exp.get("company", "")
            title = exp.get("title", "")
            start = exp.get("start_date", "")
            end = exp.get("end_date", "") or "Present"
            location = exp.get("location", "")
            bullets = _normalize_bullets(exp.get("bullets", []))

            p_co = doc.add_paragraph()
            r_co = p_co.add_run(company)
            r_co.bold = True
            r_co.font.size = Pt(11)
            r_co.font.name = "Calibri"
            _set_para_spacing(p_co, before=80, after=0)

            p_ti = doc.add_paragraph()
            r_ti = p_ti.add_run(title)
            r_ti.italic = True
            r_ti.font.size = Pt(10)
            r_ti.font.name = "Calibri"
            meta_parts = [x for x in [location, f"{start} – {end}" if start else ""] if x]
            if meta_parts:
                _gray_run(p_ti, "   " + "  ·  ".join(meta_parts))
            _set_para_spacing(p_ti, before=0, after=30)

            for bullet in bullets:
                bp = doc.add_paragraph()
                r_dot = bp.add_run("▸  ")
                r_dot.font.color.rgb = RGBColor(
                    int(_ACCENT[:2], 16), int(_ACCENT[2:4], 16), int(_ACCENT[4:], 16)
                )
                r_dot.font.size = Pt(9)
                r_dot.font.name = "Calibri"
                r_bul = bp.add_run(bullet)
                r_bul.font.size = Pt(10)
                r_bul.font.name = "Calibri"
                bp.paragraph_format.left_indent = Inches(0.15)
                _set_para_spacing(bp, before=0, after=20)

    # ── Education ──
    education = data.get("education") or []
    if education:
        section_heading("Education")
        for edu in education:
            institution = edu.get("institution", "")
            degree = edu.get("degree", "")
            field = edu.get("field", "")
            start = edu.get("start_date", "")
            end = edu.get("end_date", "")
            gpa = edu.get("gpa", "")

            p_inst = doc.add_paragraph()
            r_inst = p_inst.add_run(institution)
            r_inst.bold = True
            r_inst.font.size = Pt(11)
            r_inst.font.name = "Calibri"
            _set_para_spacing(p_inst, before=80, after=0)

            p_deg = doc.add_paragraph()
            deg_str = degree
            if field:
                deg_str += f" in {field}"
            r_deg = p_deg.add_run(deg_str)
            r_deg.italic = True
            r_deg.font.size = Pt(10)
            meta_parts = []
            if start:
                meta_parts.append(f"{start} – {end}")
            if gpa:
                meta_parts.append(f"GPA: {gpa}")
            if meta_parts:
                _gray_run(p_deg, "   " + "  ·  ".join(meta_parts))
            _set_para_spacing(p_deg, before=0, after=30)

    # ── Skills ──
    skill_categories = _normalize_skill_categories(data.get("skills"))
    if skill_categories:
        section_heading("Skills")
        for cat in skill_categories:
            p_cat = doc.add_paragraph()
            r_cat = p_cat.add_run(cat["category"] + "  ")
            r_cat.bold = True
            r_cat.font.size = Pt(10)
            r_cat.font.name = "Calibri"
            r_cat.font.color.rgb = RGBColor(
                int(_ACCENT[:2], 16), int(_ACCENT[2:4], 16), int(_ACCENT[4:], 16)
            )
            # Each skill: name + dots (colored) + label
            for i, sk in enumerate(cat["skills"]):
                if i > 0:
                    sep = p_cat.add_run("     ")
                    sep.font.size = Pt(10)
                label = _LEVEL_LABELS[max(1, min(5, sk["level"]))]
                r_name = p_cat.add_run(sk["name"] + "  ")
                r_name.font.size = Pt(10)
                r_name.font.name = "Calibri"
                r_dots = p_cat.add_run(_level_dots(sk["level"]))
                r_dots.font.size = Pt(8)
                r_dots.font.name = "Calibri"
                r_dots.font.color.rgb = RGBColor(
                    int(_ACCENT[:2], 16), int(_ACCENT[2:4], 16), int(_ACCENT[4:], 16)
                )
                r_lbl = p_cat.add_run(f"  {label}")
                r_lbl.font.size = Pt(8)
                r_lbl.font.name = "Calibri"
                r_lbl.font.color.rgb = RGBColor(
                    int(_GRAY[:2], 16), int(_GRAY[2:4], 16), int(_GRAY[4:], 16)
                )
            _set_para_spacing(p_cat, before=20, after=20)

    # ── Certifications ──
    certifications = data.get("certifications") or []
    if certifications:
        section_heading("Certifications")
        for cert in certifications:
            name = cert.get("name", "")
            issuer = cert.get("issuer", "")
            date = cert.get("date", "")
            p_cert = doc.add_paragraph()
            r_cert = p_cert.add_run(name)
            r_cert.bold = True
            r_cert.font.size = Pt(10)
            r_cert.font.name = "Calibri"
            meta = "  ·  ".join(x for x in [issuer, date] if x)
            if meta:
                _gray_run(p_cert, f"   {meta}")
            _set_para_spacing(p_cert, before=40, after=20)

    # ── Projects ──
    projects = _normalize_projects(data.get("projects", []))
    if projects:
        section_heading("Projects")
        for proj in projects:
            p_proj = doc.add_paragraph()
            r_name = p_proj.add_run(proj["name"])
            r_name.bold = True
            r_name.font.size = Pt(11)
            r_name.font.name = "Calibri"
            if proj["link"]:
                r_link = _gray_run(p_proj, f"   {proj['link']}")
                r_link.font.size = Pt(9)
            _set_para_spacing(p_proj, before=80, after=0)
            if proj["tech_stack"]:
                p_stack = doc.add_paragraph()
                r_label = p_stack.add_run("Stack  ")
                r_label.bold = True
                r_label.font.size = Pt(9)
                r_label.font.color.rgb = RGBColor(
                    int(_ACCENT[:2], 16), int(_ACCENT[2:4], 16), int(_ACCENT[4:], 16)
                )
                r_label.font.name = "Calibri"
                for i, tech in enumerate(proj["tech_stack"]):
                    if i > 0:
                        sep = p_stack.add_run("  ·  ")
                        sep.font.size = Pt(9)
                        sep.font.name = "Calibri"
                    r_tech = p_stack.add_run(tech)
                    r_tech.font.size = Pt(9)
                    r_tech.font.name = "Calibri"
                _set_para_spacing(p_stack, before=0, after=0)
            if proj["description"]:
                p_desc = doc.add_paragraph(proj["description"])
                p_desc.runs[0].font.size = Pt(10)
                p_desc.runs[0].font.name = "Calibri"
                _set_para_spacing(p_desc, before=20, after=20)

    # ── Achievements ──
    achievements = _normalize_bullets(data.get("achievements", []))
    if achievements:
        section_heading("Achievements")
        for ach in achievements:
            ap = doc.add_paragraph()
            r_dot = ap.add_run("▸  ")
            r_dot.font.color.rgb = RGBColor(
                int(_ACCENT[:2], 16), int(_ACCENT[2:4], 16), int(_ACCENT[4:], 16)
            )
            r_dot.font.size = Pt(9)
            ap.add_run(ach).font.size = Pt(10)
            ap.paragraph_format.left_indent = Inches(0.15)
            _set_para_spacing(ap, before=0, after=20)

    doc.save(output_path)
    return output_path
