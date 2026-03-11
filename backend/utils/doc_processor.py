import docx
import os
from typing import List, Dict

from utils.logger import get_logger
logger = get_logger(__name__)


def extract_text_from_docx(file_path: str) -> str:
    """Read all text from a .docx file."""
    if not os.path.exists(file_path):
        return ""
    try:
        doc = docx.Document(file_path)
        return "\n".join([p.text for p in doc.paragraphs])
    except Exception as e:
        logger.error("Error reading docx %s: %s", file_path, e, exc_info=True)
        return ""


def generate_tailored_docx(
    master_path: str, output_path: str, tailored_bullets: List[Dict[str, str]]
):
    """
    Replace 'old' bullet points with 'new' tailored versions in a .docx file.
    Note: This is a basic implementation that looks for exact paragraph matches.
    """
    if not os.path.exists(master_path):
        raise FileNotFoundError(f"Master resume not found at {master_path}")

    doc = docx.Document(master_path)

    # Create a mapping for faster lookup
    replacements = {b["old"].strip(): b["new"].strip() for b in tailored_bullets}

    # Iterate through paragraphs and replace text if a match is found
    for paragraph in doc.paragraphs:
        p_text = paragraph.text.strip()
        if p_text in replacements:
            # We replace the text but try to keep the first run's formatting if possible
            # A more robust version would handle multiple runs and formatting better.
            new_text = replacements[p_text]

            # Clear all runs and add the new text
            for run in paragraph.runs:
                run.text = ""
            paragraph.runs[0].text = new_text

    # Save the new document
    doc.save(output_path)
    return output_path


def save_text_to_docx(text: str, output_path: str):
    """Create a new .docx file from plain text."""
    doc = docx.Document()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    doc.save(output_path)
    return output_path
